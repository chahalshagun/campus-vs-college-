"""Generate a Word (.docx) template for the final dissertation.

This script builds a structured DOCX document with all required NTCC/SIL
project-report sections, using only English placeholders. It produces
`dissertation_template.docx` in this same folder.

Prerequisite:
    pip install python-docx

Usage (from project root):
    python "final disseratation/generate_dissertation_doc.py"
"""
from datetime import date
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit(
        "python-docx is not installed. Install it with 'pip install python-docx' "
        "and then re-run this script."
    ) from exc


OUTPUT_NAME = "dissertation_template.docx"


def add_heading(doc: 'Document', text: str, level: int = 1):
    """Utility: add heading and blank line."""
    doc.add_heading(text, level=level)
    doc.add_paragraph("")


def build_document() -> 'Document':
    doc = Document()

    # ---------------------- Preliminary Pages ---------------------------
    add_heading(doc, "Title Page", 0)
    doc.add_paragraph("[Project Title]")
    doc.add_paragraph("[Student Name]")
    doc.add_paragraph("[Enrollment Number]")
    doc.add_paragraph("[Institute Name]")
    doc.add_paragraph(f"Submission Date: {date.today():%B %d, %Y}")
    doc.add_page_break()

    add_heading(doc, "Declaration")
    doc.add_paragraph("[Insert declaration text here]")
    doc.add_page_break()

    add_heading(doc, "Faculty Guide Approval")
    doc.add_paragraph("[Insert approval statement/signature block]")
    doc.add_page_break()

    add_heading(doc, "Acknowledgements")
    doc.add_paragraph("[Insert acknowledgements]")
    doc.add_page_break()

    add_heading(doc, "Abstract")
    doc.add_paragraph("[Insert abstract]")
    doc.add_page_break()

    add_heading(doc, "Table of Contents")
    doc.add_paragraph("(Use References → Table of Contents in Word to generate this section once content is complete.)")
    doc.add_page_break()

    add_heading(doc, "List of Tables")
    doc.add_paragraph("(Use References → Insert Table of Figures and choose Table.)")
    doc.add_page_break()

    add_heading(doc, "List of Figures")
    doc.add_paragraph("(Use References → Insert Table of Figures.)")
    doc.add_page_break()

    # ------------------------ Main Text ---------------------------------
    for chapter in range(1, 4):
        add_heading(doc, f"Chapter {chapter}: [Title]")
        add_heading(doc, "1.1 Section Title", level=2)
        doc.add_paragraph("[Write your content here]")
        doc.add_page_break()

    # -------------------- Reference Material ---------------------------
    add_heading(doc, "References")
    doc.add_paragraph("[Insert bibliography entries]")
    doc.add_page_break()

    add_heading(doc, "Appendix A")
    doc.add_paragraph("[Appendix content]")
    doc.add_page_break()

    add_heading(doc, "Appendix B")
    doc.add_paragraph("[Appendix content]")

    return doc


def main():
    output_path = Path(__file__).with_name(OUTPUT_NAME)
    doc = build_document()
    doc.save(output_path)
    print(f"Dissertation template created at: {output_path.resolve()}")


if __name__ == "__main__":
    main()
